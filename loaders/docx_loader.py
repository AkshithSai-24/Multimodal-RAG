"""
DOCX Loader — python-docx for text/tables, raw XML for embedded images.
"""

from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path
from typing import List

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from loaders.base_loader import BaseLoader
from models.document import DocumentChunk, Modality, SourceType

# OOXML image relationship type
_IMG_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


class DOCXLoader(BaseLoader):
    """Load text, tables, and embedded images from a .docx file."""

    async def load(self, source: str, **kwargs) -> List[DocumentChunk]:
        path = Path(source)
        chunks: List[DocumentChunk] = []
        source_name = path.name
        doc = DocxDocument(str(path))

        # ── 1. Text paragraphs ────────────────────────────────────────────────
        current_block: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                if current_block:
                    chunks.append(
                        DocumentChunk(
                            content="\n".join(current_block),
                            modality=Modality.TEXT,
                            source_type=SourceType.DOCX,
                            source_id=str(path),
                            source_name=source_name,
                        )
                    )
                    current_block = []
            else:
                current_block.append(text)

        if current_block:
            chunks.append(
                DocumentChunk(
                    content="\n".join(current_block),
                    modality=Modality.TEXT,
                    source_type=SourceType.DOCX,
                    source_id=str(path),
                    source_name=source_name,
                )
            )

        # ── 2. Tables ─────────────────────────────────────────────────────────
        for tbl_idx, table in enumerate(doc.tables):
            md = self._table_to_markdown(table)
            if md:
                chunks.append(
                    DocumentChunk(
                        content=md,
                        modality=Modality.TABLE,
                        source_type=SourceType.DOCX,
                        source_id=str(path),
                        source_name=source_name,
                        metadata={"table_index": tbl_idx},
                    )
                )

        # ── 3. Embedded images (via zip + relationship parts) ─────────────────
        with zipfile.ZipFile(str(path), "r") as zf:
            names = zf.namelist()
            media_files = [n for n in names if n.startswith("word/media/")]
            for img_idx, media_path in enumerate(media_files):
                img_bytes = zf.read(media_path)
                if len(img_bytes) < 4096:
                    continue
                ext = media_path.rsplit(".", 1)[-1].lower()
                mime = f"image/{ext}" if ext in {"jpg", "jpeg", "png", "gif", "webp"} else "image/jpeg"
                b64 = base64.b64encode(img_bytes).decode()
                chunks.append(
                    DocumentChunk(
                        content=f"Embedded image {img_idx + 1} from {source_name}",
                        modality=Modality.IMAGE,
                        source_type=SourceType.DOCX,
                        source_id=str(path),
                        source_name=source_name,
                        image_base64=b64,
                        image_mime_type=mime,
                        metadata={"image_index": img_idx, "media_path": media_path},
                    )
                )

        return chunks

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def _table_to_markdown(table) -> str:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            return ""
        header = "| " + " | ".join(rows[0]) + " |"
        separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
        body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
        return "\n".join(filter(None, [header, separator, body]))
